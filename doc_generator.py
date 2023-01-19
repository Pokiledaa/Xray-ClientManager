from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import socket
from statics import Directories




class DocGenerator:
    def __init__(self):
        self.hostname = socket.gethostname()

    
    def generate_docx(self,email,uuid):
        parsed_email = email.split("@")
        client_identity = parsed_email[1]
        doc = Document()
        doc.add_heading(client_identity).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('%s           %s' %(uuid,self.hostname), style='Intense Quote').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER        
        doc.add_paragraph("VMESS TCP OBFS").paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(f"{Directories.GENERATED_OUTPUT}/{client_identity}_VMESS-TCP-OBFS.PNG",width=Inches(6)) 
        doc.add_page_break()
        doc.add_paragraph("VLESS TCP TLS").paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(f"{Directories.GENERATED_OUTPUT}/{client_identity}_VLESS-TCP-TLS.PNG",width=Inches(6))
        doc.add_page_break()
        doc.add_paragraph("VLESS WS TLS").paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(f"{Directories.GENERATED_OUTPUT}/{client_identity}_VLESS-WS-TLS.PNG",width=Inches(6))   
        doc.save(f"{Directories.GENERATED_DOCX}/{client_identity}.docx")
    

    def generate_docx_v_1_1(self,email,uuid,append_1,append_2,append_3,append_4,append_5,append_6,append_7,append_8):
        parsed_email = email.split("@")
        client_identity = parsed_email[1]
        doc = Document()
        doc.add_heading(client_identity).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('%s           %s' %(uuid,self.hostname), style='Intense Quote').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER        
        doc.add_paragraph(append_1).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(f"{Directories.GENERATED_OUTPUT}/{client_identity}_{append_1}.PNG",width=Inches(6)) 
        doc.add_page_break()
        doc.add_paragraph(append_2).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(f"{Directories.GENERATED_OUTPUT}/{client_identity}_{append_2}.PNG",width=Inches(6)) 
        doc.add_page_break()
        doc.add_paragraph(append_3).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(f"{Directories.GENERATED_OUTPUT}/{client_identity}_{append_3}.PNG",width=Inches(6)) 
        doc.add_page_break()
        doc.add_paragraph(append_4).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(f"{Directories.GENERATED_OUTPUT}/{client_identity}_{append_4}.PNG",width=Inches(6)) 
        doc.add_page_break()
        doc.add_paragraph(append_5).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(f"{Directories.GENERATED_OUTPUT}/{client_identity}_{append_5}.PNG",width=Inches(6)) 
        doc.add_page_break()
        doc.add_paragraph(append_6).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(f"{Directories.GENERATED_OUTPUT}/{client_identity}_{append_6}.PNG",width=Inches(6)) 
        doc.add_page_break()
        doc.add_paragraph(append_7).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(f"{Directories.GENERATED_OUTPUT}/{client_identity}_{append_7}.PNG",width=Inches(6)) 
        doc.add_page_break()
        doc.add_paragraph(append_8).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture(f"{Directories.GENERATED_OUTPUT}/{client_identity}_{append_8}.PNG",width=Inches(6)) 
       
        doc.save(f"{Directories.GENERATED_DOCX}/{client_identity}.docx")
    





